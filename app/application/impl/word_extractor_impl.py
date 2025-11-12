import logging
from docx import Document
import os
from app.application.file_extractor import FileExtractor

logger = logging.getLogger(__name__)


class WordExtractorImpl(FileExtractor):
    """Implementation of Word document text extractor"""
    
    async def extract_text(self, file_path: str) -> str:
        """
        Extract text from Word document (.docx only)
        
        Args:
            file_path: Path to Word file
            
        Returns:
            Extracted text
        """
        try:
            if not os.path.exists(file_path):
                raise FileNotFoundError(f"File not found: {file_path}")
            
            file_ext = os.path.splitext(file_path)[1].lower()
            
            # Only support .docx
            if file_ext != '.docx':
                raise ValueError(
                    f"Faqat .docx formatdagi fayllar qo'llab-quvvatlanadi. "
                    f"Iltimos, .doc faylni .docx ga o'giring: "
                    f"Microsoft Word da ochib → File → Save As → .docx formatda saqlang"
                )
            
            doc = Document(file_path)
            
            # Extract text from paragraphs
            text_parts = []
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_parts.append(paragraph.text.strip())
            
            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            text_parts.append(cell.text.strip())
            
            full_text = "\n".join(text_parts)
            
            if not full_text.strip():
                raise ValueError("Hujjatda matn topilmadi")
            
            logger.info(f"Extracted {len(full_text)} characters from {file_path}")
            return full_text
            
        except Exception as e:
            logger.error(f"Error extracting text from {file_path}: {str(e)}")
            raise